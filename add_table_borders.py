#!/usr/bin/env python3
"""Add borders to tables in Word documents."""

import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_table_borders(table):
    """Add black borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

    # Define border style
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def add_borders_to_docx(filepath):
    """Add borders to all tables in a Word document."""
    doc = Document(filepath)

    for table in doc.tables:
        set_table_borders(table)

    doc.save(filepath)
    print(f"Updated: {filepath}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python add_table_borders.py <file.docx> [file2.docx ...]")
        sys.exit(1)

    for filepath in sys.argv[1:]:
        try:
            add_borders_to_docx(filepath)
        except Exception as e:
            print(f"Error processing {filepath}: {e}")
