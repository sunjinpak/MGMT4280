#!/usr/bin/env python3
"""Convert Syllabus markdown to Word with proper table formatting."""

import re
import subprocess
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_table_borders(table):
    """Add black borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def preprocess_markdown(input_file, output_file):
    """Preprocess markdown to handle line breaks in tables properly."""
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove the div wrapper for black-border-table (not needed for Word)
    content = re.sub(r'<div class="black-border-table" markdown="1">\s*\n?', '', content)
    content = re.sub(r'\n?</div>', '', content)

    # Remove Jekyll front matter
    content = re.sub(r'^---\n.*?\n---\n', '', content, flags=re.DOTALL)

    # Remove navigation links at top
    content = re.sub(r'^\[Home\]\(index\)\n+', '', content)
    content = re.sub(r'^\*\*\[Download Word Document\].*?\*\*\n+', '', content)

    # Convert relative links to absolute GitHub Pages URLs
    base_url = "https://sunjinpak.github.io/MGMT4280/"

    # Find all markdown links and convert relative ones to absolute
    def convert_link(match):
        text = match.group(1)
        url = match.group(2)
        # Skip if already absolute URL or anchor link
        if url.startswith('http') or url.startswith('#') or url.startswith('mailto:'):
            return match.group(0)
        # Skip Word document download link
        if url.endswith('.docx'):
            return match.group(0)
        # Convert relative path to absolute
        return f'[{text}]({base_url}{url})'

    content = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', convert_link, content)

    # In table cells, replace <br> with a special marker that will survive pandoc
    # Use Unicode line separator character
    content = content.replace('<br>', '⏎')

    # Ensure bullet lists have a blank line before them (pandoc requirement)
    # Pattern: non-blank line ending with text, followed by a line starting with "- "
    content = re.sub(r'([^\n])\n(- )', r'\1\n\n\2', content)

    # Replace horizontal rules within table cells (-----) with newline marker
    # This pattern matches table cells containing -----
    def replace_dashes_in_tables(match):
        cell_content = match.group(0)
        # Replace ----- patterns with nothing (they're just visual separators)
        cell_content = re.sub(r'-{5,}', '', cell_content)
        return cell_content

    # Process each table row
    lines = content.split('\n')
    processed_lines = []
    for line in lines:
        if line.startswith('|') and '-----' in line and not re.match(r'^\|[-:\s|]+\|$', line):
            # This is a table data row with dashes - remove them
            line = re.sub(r'-{5,}', '', line)
        processed_lines.append(line)
    content = '\n'.join(processed_lines)

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(content)

def escape_xml(text):
    """Escape special XML characters."""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&apos;')
    return text

def convert_markers_to_linebreaks_xml(cell):
    """Convert ⏎ markers to actual line breaks in a cell, preserving hyperlinks."""
    full_text = cell.text

    if '⏎' not in full_text:
        return False

    # Find all text elements (w:t) in the cell and replace marker with line break
    tc = cell._tc
    text_elements = tc.findall('.//' + qn('w:t'))

    for t_elem in text_elements:
        if t_elem.text and '⏎' in t_elem.text:
            parts = t_elem.text.split('⏎')
            if len(parts) > 1:
                # Get the parent run (w:r)
                parent_r = t_elem.getparent()

                # Set first part as the text
                t_elem.text = parts[0].strip()

                # Insert line breaks and text for remaining parts
                insert_after = t_elem
                for part in parts[1:]:
                    part = part.strip()
                    # Create line break element
                    br = parse_xml(r'<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    insert_after.addnext(br)
                    insert_after = br

                    if part:
                        # Create new text element with escaped text
                        escaped_part = escape_xml(part)
                        new_t = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve">{escaped_part}</w:t>')
                        insert_after.addnext(new_t)
                        insert_after = new_t

    return True

def postprocess_word(docx_file):
    """Post-process Word document to fix table formatting."""
    doc = Document(docx_file)

    # Add borders to all tables
    for table in doc.tables:
        set_table_borders(table)

    # Process ALL tables for line break markers
    for table_idx, table in enumerate(doc.tables):
        fixed_cells = 0

        for row in table.rows:
            for cell in row.cells:
                if convert_markers_to_linebreaks_xml(cell):
                    fixed_cells += 1

        if fixed_cells > 0:
            print(f"Table {table_idx}: Fixed {fixed_cells} cells with line breaks")

    # Set column widths for specific tables
    for table in doc.tables:
        first_cell = table.rows[0].cells[0].text.strip().lower()

        # Weekly Schedule table (2 columns: Class, Contents)
        if len(table.columns) == 2 and 'class' in first_cell:
            print(f"Setting column widths for Weekly Schedule table")
            for row in table.rows:
                row.cells[0].width = Cm(3)
                row.cells[1].width = Cm(15)

        # Grading Criteria table (3 columns: Assessment, Points, %)
        if len(table.columns) == 3 and 'assessment' in first_cell:
            print(f"Setting column widths for Grading Criteria table")
            for row in table.rows:
                row.cells[0].width = Cm(9)   # Assessment
                row.cells[1].width = Cm(5)   # Points
                row.cells[2].width = Cm(2)   # %

    # Make all hyperlinks blue and underlined by modifying XML directly
    hyperlink_count = 0

    # Find all w:hyperlink elements in the document body
    body = doc._body._body
    hyperlinks = body.findall('.//' + qn('w:hyperlink'))

    for hyperlink in hyperlinks:
        # Find all runs inside this hyperlink
        runs = hyperlink.findall('.//' + qn('w:r'))
        for run_elem in runs:
            # Get or create rPr (run properties)
            rPr = run_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                run_elem.insert(0, rPr)

            # Add blue color
            color = rPr.find(qn('w:color'))
            if color is None:
                color = parse_xml(r'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="0066CC"/>')
                rPr.append(color)
            else:
                color.set(qn('w:val'), '0066CC')

            # Add underline
            u = rPr.find(qn('w:u'))
            if u is None:
                u = parse_xml(r'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
                rPr.append(u)

            hyperlink_count += 1

    print(f"Made {hyperlink_count} hyperlink runs blue")

    doc.save(docx_file)
    print(f"Saved: {docx_file}")

def main():
    os.chdir('/tmp/MGMT4280')

    # Step 1: Preprocess markdown
    print("Step 1: Preprocessing markdown...")
    preprocess_markdown('syllabus.md', 'syllabus_temp.md')

    # Step 2: Convert to Word with pandoc
    print("Step 2: Converting to Word with pandoc...")
    subprocess.run([
        'pandoc', 'syllabus_temp.md',
        '-o', 'Syllabus.docx',
        '--from=markdown+raw_html'
    ], check=True)

    # Step 3: Post-process Word document
    print("Step 3: Post-processing Word document...")
    postprocess_word('Syllabus.docx')

    # Cleanup
    os.remove('syllabus_temp.md')
    print("Done!")

if __name__ == "__main__":
    main()
