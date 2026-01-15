#!/usr/bin/env python3
"""Fix table column widths and line breaks in Syllabus Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK
import re

def set_table_borders(table):
    """Add black borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def fix_cell_line_breaks(cell):
    """Replace <br> tags with actual line breaks in a cell."""
    # Get all text from the cell
    full_text = cell.text

    # Check if there are <br> or <br/> tags
    if '<br>' in full_text or '<br/>' in full_text:
        # Clear the cell
        cell.text = ""

        # Split by <br> tags
        parts = re.split(r'<br\s*/?>', full_text)

        # Add each part as a new paragraph or with line break
        paragraph = cell.paragraphs[0]
        for i, part in enumerate(parts):
            if i > 0:
                # Add line break before subsequent parts
                run = paragraph.add_run()
                run.add_break()
            paragraph.add_run(part.strip())

def fix_syllabus():
    """Fix the Syllabus Word document."""
    doc = Document("Syllabus.docx")

    for table in doc.tables:
        # Add borders to all tables
        set_table_borders(table)

        # Check if this is the Weekly Schedule table (2 columns: Class, Contents)
        if len(table.columns) == 2:
            # Check first row header
            first_cell_text = table.rows[0].cells[0].text.strip().lower()
            if 'class' in first_cell_text or 'week' in first_cell_text:
                print(f"Found Weekly Schedule table with {len(table.rows)} rows")

                # Set column widths: Class = 2cm, Contents = rest
                for row in table.rows:
                    row.cells[0].width = Cm(3)  # About 1.2 inches
                    row.cells[1].width = Cm(15)  # About 6 inches

                    # Fix line breaks in the contents cell
                    fix_cell_line_breaks(row.cells[1])

                print("Fixed Weekly Schedule table")

    doc.save("Syllabus.docx")
    print("Saved: Syllabus.docx")

if __name__ == "__main__":
    fix_syllabus()
